"""P2 #10 — Executive EVM + DCMA report builder (DOCX + PDF).

Pure builder: takes a structured `data` dict (project meta + EVM metrics +
forecast + DCMA rules + finish-driver) and writes a branded one/two-page
executive report. No COM, no MS Project — fully testable.

Brand (Mühendis İnşaat Yönetim A.Ş. — feedback_brand_muhendis_not_mcs):
    LACIVERT #0B1F4D · NAVY #3D4663 · TURQUOISE #39B4CC · LABEL_GRAY #6B7394

`data` shape (all keys optional, builder is defensive):
{
  "project":  {"name", "file", "status_date", "forecast_finish"},
  "evm":      {"bac","pv","ev","ac","spi","cpi","sv","cv",
               "eac_t2","vac","rag","completion_pct","currency_mode"},
  "dcma":     {"rules":[{id,name,threshold,actual,actual_unit,status}],
               "summary":{pass_count,fail_count,overall_rag}},
  "driver":   {"anomaly","gap_days","driver":{...}} | None,
  "generated_for": "Müşteri/Proje",
}
"""
from typing import Dict, Any, List

COMPANY = "Mühendis İnşaat Yönetim A.Ş."
LACIVERT = "0B1F4D"
NAVY = "3D4663"
TURQUOISE = "39B4CC"
LABEL_GRAY = "6B7394"
ZEBRA = "F0F3F8"

_RAG_HEX = {"green": "2E7D32", "amber": "E68A00", "red": "C62828"}


def _fmt(v, nd=1):
    if v is None:
        return "—"
    try:
        return f"{float(v):,.{nd}f}"
    except (ValueError, TypeError):
        return str(v)


def _evm_rows(evm: Dict[str, Any]) -> List[List[str]]:
    unit = "$" if (evm.get("currency_mode") == "cost") else "saat"
    return [
        ["BAC (Bütçe)", f"{_fmt(evm.get('bac'))} {unit}"],
        ["PV (Planlanan Değer)", f"{_fmt(evm.get('pv'))} {unit}"],
        ["EV (Kazanılan Değer)", f"{_fmt(evm.get('ev'))} {unit}"],
        ["AC (Gerçekleşen Maliyet)", f"{_fmt(evm.get('ac'))} {unit}"],
        ["SPI (Program Endeksi)", _fmt(evm.get("spi"), 3)],
        ["CPI (Maliyet Endeksi)", _fmt(evm.get("cpi"), 3)],
        ["SV (Program Sapması)", f"{_fmt(evm.get('sv'))} {unit}"],
        ["CV (Maliyet Sapması)", f"{_fmt(evm.get('cv'))} {unit}"],
        ["EAC (Tahmini Toplam)", f"{_fmt(evm.get('eac_t2'))} {unit}"],
        ["VAC (Tahmini Sapma)", f"{_fmt(evm.get('vac'))} {unit}"],
        ["Tamamlanma %", f"{_fmt(evm.get('completion_pct'))} %"],
    ]


# ============================================================================
# DOCX
# ============================================================================

def build_executive_docx(data: Dict[str, Any], output_path: str) -> str:
    """Build the executive report as a .docx. Returns output_path."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    def _rgb(hexs):
        return RGBColor(int(hexs[0:2], 16), int(hexs[2:4], 16), int(hexs[4:6], 16))

    def _shade(cell, hexs):
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), hexs)
        tcPr.append(shd)

    proj = data.get("project", {}) or {}
    evm = data.get("evm", {}) or {}
    dcma = data.get("dcma", {}) or {}
    driver = data.get("driver")

    doc = Document()
    for s in doc.sections:
        s.top_margin = Mm(16); s.bottom_margin = Mm(16)
        s.left_margin = Mm(16); s.right_margin = Mm(16)

    # --- Title band ---
    h = doc.add_paragraph()
    r = h.add_run(COMPANY)
    r.font.name = "Calibri"; r.font.size = Pt(11); r.font.bold = True
    r.font.color.rgb = _rgb(LABEL_GRAY)

    t = doc.add_paragraph()
    rt = t.add_run("EVM & DCMA Yönetici Raporu")
    rt.font.name = "Calibri"; rt.font.size = Pt(22); rt.font.bold = True
    rt.font.color.rgb = _rgb(LACIVERT)

    meta = doc.add_paragraph()
    rm = meta.add_run(
        f"Proje: {proj.get('name','—')}    ·    Veri Tarihi: "
        f"{proj.get('status_date','—')}    ·    Forecast Finish: "
        f"{proj.get('forecast_finish','—')}")
    rm.font.name = "Calibri"; rm.font.size = Pt(10); rm.font.color.rgb = _rgb(NAVY)

    # --- RAG banner ---
    rag = (evm.get("rag") or dcma.get("summary", {}).get("overall_rag") or "—")
    rb = doc.add_paragraph()
    rbr = rb.add_run(f"Genel Durum (RAG): {str(rag).upper()}")
    rbr.font.name = "Calibri"; rbr.font.size = Pt(12); rbr.font.bold = True
    rbr.font.color.rgb = _rgb(_RAG_HEX.get(str(rag).lower(), NAVY))

    # --- EVM table ---
    doc.add_paragraph().add_run("1. Earned Value (PMI PMBOK 8th)").bold = True
    rows = _evm_rows(evm)
    tbl = doc.add_table(rows=len(rows) + 1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = tbl.rows[0].cells
    for i, txt in enumerate(("Metrik", "Değer")):
        hdr[i].text = ""
        rr = hdr[i].paragraphs[0].add_run(txt)
        rr.font.bold = True; rr.font.color.rgb = _rgb("FFFFFF")
        rr.font.name = "Calibri"; rr.font.size = Pt(10)
        _shade(hdr[i], NAVY)
    for ri, (k, v) in enumerate(rows, start=1):
        c = tbl.rows[ri].cells
        c[0].text = k; c[1].text = v
        if ri % 2 == 0:
            _shade(c[0], ZEBRA); _shade(c[1], ZEBRA)

    # --- DCMA table ---
    doc.add_paragraph()
    dsum = dcma.get("summary", {}) or {}
    p = doc.add_paragraph()
    p.add_run(f"2. DCMA 14-Point Schedule Health  "
              f"({dsum.get('pass_count','?')}/14 PASS · "
              f"{str(dsum.get('overall_rag','—')).upper()})").bold = True
    drules = dcma.get("rules", []) or []
    dt = doc.add_table(rows=len(drules) + 1, cols=4)
    dhdr = dt.rows[0].cells
    for i, txt in enumerate(("#  Kural", "Eşik", "Gerçekleşen", "Durum")):
        dhdr[i].text = ""
        rr = dhdr[i].paragraphs[0].add_run(txt)
        rr.font.bold = True; rr.font.color.rgb = _rgb("FFFFFF")
        rr.font.name = "Calibri"; rr.font.size = Pt(9)
        _shade(dhdr[i], NAVY)
    for ri, rule in enumerate(drules, start=1):
        c = dt.rows[ri].cells
        c[0].text = f"{rule.get('id','')}  {rule.get('name','')}"
        c[1].text = str(rule.get("threshold", ""))
        c[2].text = f"{rule.get('actual','')}{rule.get('actual_unit','')}"
        st = str(rule.get("status", "")).upper()
        c[3].text = ""
        sr = c[3].paragraphs[0].add_run(st)
        sr.font.bold = True
        sr.font.color.rgb = _rgb("2E7D32" if st == "PASS" else "C62828")
        for cc in c:
            for pp in cc.paragraphs:
                for rn in pp.runs:
                    rn.font.size = Pt(9); rn.font.name = "Calibri"
        if ri % 2 == 0:
            for cc in c:
                _shade(cc, ZEBRA)

    # --- Forecast-driver finding ---
    if driver and driver.get("anomaly"):
        doc.add_paragraph()
        doc.add_paragraph().add_run(
            "3. Kritik Bulgu — Forecast Finish Sürücüsü (RULE 16.C)").bold = True
        d = driver.get("driver", {}) or {}
        note = d.get("note", "")
        fp = doc.add_paragraph()
        fr = fp.add_run(
            f"Anomali: '{d.get('wbs_name','?')}' branch'i diğerlerinden "
            f"{driver.get('gap_days','?')} gün geç bitiyor. "
            f"Sürücü aktivite: {d.get('driving_task',{}).get('code','?')} "
            f"({d.get('driving_task',{}).get('task_type','?')}). {note}")
        fr.font.name = "Calibri"; fr.font.size = Pt(10)
        fr.font.color.rgb = _rgb("C62828")

    # --- Footer line ---
    doc.add_paragraph()
    f = doc.add_paragraph()
    fr = f.add_run(f"{COMPANY} · PMI PMBOK 8th · DCMA 14-Point · "
                   "Industry Standard")
    fr.font.name = "Calibri"; fr.font.size = Pt(8)
    fr.font.color.rgb = _rgb(LABEL_GRAY)
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(output_path)
    return output_path


# ============================================================================
# PDF
# ============================================================================

def build_executive_pdf(data: Dict[str, Any], output_path: str) -> str:
    """Build the executive report as a .pdf via reportlab. Returns output_path."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                    TableStyle)
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

    proj = data.get("project", {}) or {}
    evm = data.get("evm", {}) or {}
    dcma = data.get("dcma", {}) or {}
    driver = data.get("driver")

    lac = colors.HexColor("#0B1F4D")
    navy = colors.HexColor("#3D4663")
    gray = colors.HexColor("#6B7394")
    zebra = colors.HexColor("#F0F3F8")

    styles = getSampleStyleSheet()
    h_title = ParagraphStyle("t", parent=styles["Title"], textColor=lac,
                             fontSize=20, alignment=0)
    h_sub = ParagraphStyle("s", parent=styles["Normal"], textColor=navy,
                           fontSize=10)
    h_sec = ParagraphStyle("sec", parent=styles["Heading2"], textColor=lac,
                           fontSize=12)
    small = ParagraphStyle("sm", parent=styles["Normal"], textColor=gray,
                           fontSize=8, alignment=1)

    doc = SimpleDocTemplate(output_path, pagesize=A4,
                            topMargin=16 * mm, bottomMargin=16 * mm,
                            leftMargin=16 * mm, rightMargin=16 * mm)
    el = []
    el.append(Paragraph(COMPANY, ParagraphStyle("c", parent=styles["Normal"],
              textColor=gray, fontSize=10)))
    el.append(Paragraph("EVM &amp; DCMA Yönetici Raporu", h_title))
    el.append(Paragraph(
        f"Proje: {proj.get('name','—')} &nbsp;·&nbsp; Veri Tarihi: "
        f"{proj.get('status_date','—')} &nbsp;·&nbsp; Forecast Finish: "
        f"{proj.get('forecast_finish','—')}", h_sub))
    rag = (evm.get("rag") or dcma.get("summary", {}).get("overall_rag") or "—")
    el.append(Paragraph(
        f"<b>Genel Durum (RAG): {str(rag).upper()}</b>",
        ParagraphStyle("rag", parent=styles["Normal"], fontSize=12,
                       textColor=colors.HexColor("#" + _RAG_HEX.get(
                           str(rag).lower(), NAVY)))))
    el.append(Spacer(1, 6))

    # EVM table
    el.append(Paragraph("1. Earned Value (PMI PMBOK 8th)", h_sec))
    edata = [["Metrik", "Değer"]] + _evm_rows(evm)
    et = Table(edata, colWidths=[90 * mm, 80 * mm])
    est = [("BACKGROUND", (0, 0), (-1, 0), navy),
           ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
           ("FONTSIZE", (0, 0), (-1, -1), 9),
           ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#E3E7F0"))]
    for ri in range(1, len(edata)):
        if ri % 2 == 0:
            est.append(("BACKGROUND", (0, ri), (-1, ri), zebra))
    et.setStyle(TableStyle(est))
    el.append(et)
    el.append(Spacer(1, 8))

    # DCMA table
    dsum = dcma.get("summary", {}) or {}
    el.append(Paragraph(
        f"2. DCMA 14-Point ({dsum.get('pass_count','?')}/14 PASS · "
        f"{str(dsum.get('overall_rag','—')).upper()})", h_sec))
    ddata = [["#  Kural", "Eşik", "Gerçekleşen", "Durum"]]
    drules = dcma.get("rules", []) or []
    for rule in drules:
        ddata.append([
            f"{rule.get('id','')}  {rule.get('name','')}",
            str(rule.get("threshold", "")),
            f"{rule.get('actual','')}{rule.get('actual_unit','')}",
            str(rule.get("status", "")).upper()])
    dtab = Table(ddata, colWidths=[80 * mm, 25 * mm, 35 * mm, 30 * mm])
    dst = [("BACKGROUND", (0, 0), (-1, 0), navy),
           ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
           ("FONTSIZE", (0, 0), (-1, -1), 8),
           ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#E3E7F0"))]
    for ri in range(1, len(ddata)):
        st = ddata[ri][3]
        dst.append(("TEXTCOLOR", (3, ri), (3, ri),
                    colors.HexColor("#2E7D32" if st == "PASS" else "#C62828")))
        if ri % 2 == 0:
            dst.append(("BACKGROUND", (0, ri), (-1, ri), zebra))
    dtab.setStyle(TableStyle(dst))
    el.append(dtab)

    if driver and driver.get("anomaly"):
        d = driver.get("driver", {}) or {}
        el.append(Spacer(1, 8))
        el.append(Paragraph("3. Kritik Bulgu — Forecast Driver (RULE 16.C)",
                            h_sec))
        el.append(Paragraph(
            f"<font color='#C62828'>Anomali: '{d.get('wbs_name','?')}' "
            f"{driver.get('gap_days','?')} gün geç bitiyor. Sürücü: "
            f"{d.get('driving_task',{}).get('code','?')} "
            f"({d.get('driving_task',{}).get('task_type','?')}). "
            f"{d.get('note','')}</font>", styles["Normal"]))

    el.append(Spacer(1, 10))
    el.append(Paragraph(
        f"{COMPANY} · PMI PMBOK 8th · DCMA 14-Point · Industry Standard",
        small))
    doc.build(el)
    return output_path
