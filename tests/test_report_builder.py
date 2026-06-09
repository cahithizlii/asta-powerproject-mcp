"""P2 #10 — executive report builder (DOCX + PDF)."""
import os
import pytest

from report_builder import build_executive_docx, build_executive_pdf, COMPANY

SAMPLE = {
    "project": {"name": "CAU Hospital", "file": "cau.xer",
                "status_date": "2026-05-01", "forecast_finish": "2027-01-11"},
    "evm": {"bac": 2505038.0, "pv": 1200000.0, "ev": 1100000.0,
            "ac": 1399158.0, "spi": 0.92, "cpi": 0.79, "sv": -100000.0,
            "cv": -299158.0, "eac_t2": 3170000.0, "vac": -665000.0,
            "rag": "amber", "completion_pct": 43.9, "currency_mode": "hours"},
    "dcma": {"rules": [
        {"id": i, "name": f"Rule {i}", "threshold": "<5%", "actual": 3.0,
         "actual_unit": "%", "status": "pass" if i % 2 else "fail"}
        for i in range(1, 15)],
        "summary": {"pass_count": 7, "fail_count": 7, "overall_rag": "amber"}},
    "driver": {"anomaly": True, "gap_days": 222, "driver": {
        "wbs_name": "Procurement", "is_loe": True,
        "driving_task": {"code": "PR-HG-ELEC-FD-3450", "task_type": "TT_LOE"},
        "note": "LOE drags finish 5 months."}},
}


def test_build_docx_creates_file(tmp_path):
    out = str(tmp_path / "exec.docx")
    p = build_executive_docx(SAMPLE, out)
    assert os.path.exists(p)
    assert os.path.getsize(p) > 5000


def test_docx_contains_key_text(tmp_path):
    from docx import Document
    out = str(tmp_path / "exec2.docx")
    build_executive_docx(SAMPLE, out)
    doc = Document(out)
    text = "\n".join(par.text for par in doc.paragraphs)
    table_text = " ".join(
        c.text for tb in doc.tables for row in tb.rows for c in row.cells)
    blob = text + " " + table_text
    assert COMPANY in blob
    assert "DCMA" in blob
    assert "SPI" in blob
    assert "Procurement" in blob          # driver finding rendered
    assert "PR-HG-ELEC-FD-3450" in blob


def test_build_pdf_creates_valid_file(tmp_path):
    out = str(tmp_path / "exec.pdf")
    p = build_executive_pdf(SAMPLE, out)
    assert os.path.exists(p)
    with open(p, "rb") as f:
        head = f.read(5)
    assert head == b"%PDF-"
    assert os.path.getsize(p) > 2000


def test_builder_defensive_on_empty_data(tmp_path):
    out = str(tmp_path / "empty.docx")
    build_executive_docx({}, out)
    assert os.path.exists(out)
    outp = str(tmp_path / "empty.pdf")
    build_executive_pdf({}, outp)
    assert os.path.exists(outp)


# ---------- action wired through XER ----------

def test_report_action_docx_from_xer(sample_cau_xer, tmp_path):
    from msproject_mcp_core import _msp_report_executive
    out = str(tmp_path / "cau_exec.docx")
    r = _msp_report_executive(file_path=sample_cau_xer, output_path=out,
                              fmt="docx")
    assert r["status"] == "ok"
    assert os.path.exists(out)
    assert r["dcma_pass"] is not None


def test_report_action_pdf_from_xer(sample_cau_xer, tmp_path):
    from msproject_mcp_core import _msp_report_executive
    out = str(tmp_path / "cau_exec.pdf")
    r = _msp_report_executive(file_path=sample_cau_xer, output_path=out,
                              fmt="pdf")
    assert r["status"] == "ok"
    with open(out, "rb") as f:
        assert f.read(5) == b"%PDF-"


def test_report_action_requires_output_path(sample_cau_xer):
    from msproject_mcp_core import _msp_report_executive
    r = _msp_report_executive(file_path=sample_cau_xer, output_path=None)
    assert r["status"] == "error"


def test_report_action_bad_fmt(sample_cau_xer, tmp_path):
    from msproject_mcp_core import _msp_report_executive
    r = _msp_report_executive(file_path=sample_cau_xer,
                              output_path=str(tmp_path / "x.txt"), fmt="txt")
    assert r["status"] == "error"
